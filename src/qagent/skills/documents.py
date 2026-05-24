from __future__ import annotations

import html
import re
import unicodedata
from datetime import datetime
from pathlib import Path
from typing import Any

_STORAGE_ROOT = Path.home() / "qagent-data" / "documents"
_MAX_TEXT = 32_000
_MAX_LIST = 50
_PREVIEW_CHARS = 80
_VALID_FORMATS = {"markdown", "md", "docx", "html"}


def _ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def _slugify(text: str) -> str:
    normalized = unicodedata.normalize("NFKD", text)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = ascii_text.lower()
    ascii_text = re.sub(r"[^a-z0-9]+", "-", ascii_text).strip("-")
    return ascii_text or "untitled"


def _timestamp() -> str:
    return datetime.now().strftime("%Y%m%d-%H%M%S")


def _truncate(text: str, limit: int = _MAX_TEXT) -> str:
    if len(text) <= limit:
        return text
    return text[:limit] + f"\n... [{len(text) - limit} chars truncated]"


def _find_document(name: str) -> Path | None:
    if not _STORAGE_ROOT.exists():
        return None
    candidate = _STORAGE_ROOT / name
    if candidate.is_file():
        return candidate
    slug = _slugify(name)
    matches = sorted(
        (
            p
            for p in _STORAGE_ROOT.iterdir()
            if p.is_file() and (slug in p.stem or name in p.name)
        ),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    return matches[0] if matches else None


def _write_markdown(path: Path, title: str, content: str) -> None:
    body = f"# {title}\n\n{content}\n"
    path.write_text(body, encoding="utf-8")


def _write_html(path: Path, title: str, content: str) -> None:
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    if not paragraphs and content.strip():
        paragraphs = [content.strip()]
    body_parts = [f"<p>{html.escape(p)}</p>" for p in paragraphs]
    body_html = "".join(body_parts)
    doc = (
        "<!doctype html><html><body>"
        f"<h1>{html.escape(title)}</h1>"
        f"<div>{body_html}</div>"
        "</body></html>"
    )
    path.write_text(doc, encoding="utf-8")


def _write_docx(path: Path, title: str, content: str) -> str | None:
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    document = Document()
    document.add_heading(title, level=1)
    for line in content.splitlines():
        stripped = line.rstrip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("### "):
            document.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)
    document.save(str(path))
    return None


def make_document(title: str, content: str, format: str = "markdown") -> str:
    try:
        fmt = (format or "markdown").lower().strip()
        if fmt not in _VALID_FORMATS:
            return (
                f"Error: unsupported format '{format}'. "
                "Use one of: markdown, md, docx, html"
            )

        _ensure_dir(_STORAGE_ROOT)
        slug = _slugify(title)
        stem = f"{_timestamp()}-{slug}"

        if fmt in {"markdown", "md"}:
            target = _STORAGE_ROOT / f"{stem}.md"
            _write_markdown(target, title, content)
        elif fmt == "html":
            target = _STORAGE_ROOT / f"{stem}.html"
            _write_html(target, title, content)
        else:
            target = _STORAGE_ROOT / f"{stem}.docx"
            err = _write_docx(target, title, content)
            if err is not None:
                return err

        return f"Saved document to {target}"
    except OSError as exc:
        return f"Error: failed to save document: {exc}"
    except Exception as exc:
        return f"Error: failed to save document: {exc}"


def list_documents() -> str:
    try:
        if not _STORAGE_ROOT.exists():
            return "(no documents)"
        files = [p for p in _STORAGE_ROOT.iterdir() if p.is_file()]
        if not files:
            return "(no documents)"
        files.sort(key=lambda p: p.stat().st_mtime, reverse=True)

        lines: list[str] = []
        for path in files[:_MAX_LIST]:
            try:
                size = path.stat().st_size
            except OSError:
                size = 0
            first_line = _first_body_line(path)
            lines.append(f"{path.name} | {size}B | {first_line}")
        return "\n".join(lines)
    except OSError as exc:
        return f"Error: failed to list documents: {exc}"


def _first_body_line(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        if suffix in {".md", ".html", ".htm", ".txt"}:
            text = path.read_text(encoding="utf-8", errors="replace")
            if suffix in {".html", ".htm"}:
                text = re.sub(r"<[^>]+>", " ", text)
            for line in text.splitlines():
                clean = line.strip()
                if not clean:
                    continue
                if clean.startswith("#"):
                    continue
                return clean[:_PREVIEW_CHARS]
            return ""
        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError:
                return "(docx)"
            try:
                document = Document(str(path))
                for para in document.paragraphs:
                    text = (para.text or "").strip()
                    if text:
                        return text[:_PREVIEW_CHARS]
            except Exception:
                return "(docx)"
            return ""
    except OSError:
        return ""
    return ""


def read_document(name: str) -> str:
    try:
        target = _find_document(name)
        if target is None:
            return f"Error: document not found: {name}"
        suffix = target.suffix.lower()
        if suffix == ".docx":
            try:
                from docx import Document
            except ImportError:
                return (
                    "Error: python-docx not installed. "
                    "Run: pip install python-docx"
                )
            try:
                document = Document(str(target))
                text = "\n".join(p.text for p in document.paragraphs)
            except Exception as exc:
                return f"Error: failed to read docx: {exc}"
            return _truncate(text)
        return _truncate(target.read_text(encoding="utf-8", errors="replace"))
    except OSError as exc:
        return f"Error: failed to read document: {exc}"


SCHEMAS: list[dict[str, Any]] = [
    {
        "type": "function",
        "function": {
            "name": "make_document",
            "description": (
                "Create a document (markdown, docx, or html) in the user's "
                "document store under ~/qagent-data/documents/."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "title": {
                        "type": "string",
                        "description": "Document title used for filename and main heading.",
                    },
                    "content": {
                        "type": "string",
                        "description": (
                            "Document body. For docx, lines starting with '## ' "
                            "or '### ' become headings and '- ' lines become bullets."
                        ),
                    },
                    "format": {
                        "type": "string",
                        "description": "Output format.",
                        "enum": ["markdown", "md", "docx", "html"],
                    },
                },
                "required": ["title", "content"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "list_documents",
            "description": "List saved documents (newest first) with size and first body line.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_document",
            "description": "Read a saved document by filename or slug fragment.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {
                        "type": "string",
                        "description": "Document filename or slug fragment.",
                    }
                },
                "required": ["name"],
            },
        },
    },
]


def dispatch(name: str, args: dict) -> str | None:
    if not isinstance(args, dict):
        args = {}

    if name == "make_document":
        title = args.get("title")
        content = args.get("content")
        fmt = args.get("format", "markdown")
        if not isinstance(title, str) or not isinstance(content, str):
            return "Error: make_document requires string title and content"
        if not isinstance(fmt, str):
            fmt = "markdown"
        return make_document(title, content, fmt)

    if name == "list_documents":
        return list_documents()

    if name == "read_document":
        doc_name = args.get("name")
        if not isinstance(doc_name, str):
            return "Error: read_document requires string name"
        return read_document(doc_name)

    return None
